from datetime import datetime

from docx import Document
# from docx.section import Section, _Header
# from docx.shared import Pt, RGBColor, Cm
# from docx.styles.style import _TableStyle, _ParagraphStyle
from docx.table import _Column, _Cell, _Row, Table
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
# from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from docx.text.paragraph import Paragraph

from pymorphy2 import MorphAnalyzer

from app.database import db
from app.models import student


morph = MorphAnalyzer()

wordDoc: Document = Document('Списки_классов\\Сведения 11 Н класса.docx')
t = wordDoc.tables[0]
p = wordDoc.paragraphs
# print(*[_.text for _ in p[2:-1]], sep="\n")
# assert int(input("0/1: ")), ValueError("Отмена")
cols = list(t.columns)
for num, col in enumerate(cols):
    cols[num] = col.cells

for col, par, pho, birth, bc, addr, emails in list(zip(*cols[1:]))[1:]:
    lastname, name, *patronymic = col.text.split()
    parents = par.text.split("\n")
    phones = pho.text.\
        strip(" ./:,").\
        replace("(", " ").replace(")", " ").\
        replace(".", "").\
        split("\n")
    emails: _Cell
    mother: None | list = None
    father: None | list = None
    info: dict = dict()
    info["birth_certificate"] = bc.text.strip()
    info["address"] = addr.text.strip()
    info["parent_emails"] = list(map(
        lambda x: x.strip(), emails.text.strip().split("\n")
    ))
    for parent in parents:
        if not parent:
            continue
        parent = parent.split()
        p_lastname, p_name, *p_patronymic = parent
        m_name = morph.parse(p_name)[0]
        gender = m_name.tag.gender
        case = m_name.tag.case
        if not p_patronymic:
            parent.append(None)
        if gender == "masc" and case == "nomn":
            exec(f"father = parent")
        else:
            exec(f"mother = parent")

    for phone in phones:
        if "п" in phone.lower():
            father.append(phone.replace("п", "").replace("\xa0", " ").strip())
        elif "м" in phone.lower():
            mother.append(phone.replace("м", "").replace("\xa0", " ").strip())
        elif "д" in phone.lower():
            info["home_phone"] = phone.replace("д", "").replace("\xa0", " ").strip()

    s = student.Students(
        name=name,
        lastname=lastname,
        birthdate=datetime.strptime(birth.text, "%d.%m.%Y").timestamp(),
        admission_year=2012,
        classroom_letter=0,
    )

    if patronymic:
        s.patronymic = patronymic[0]

    db.session.add(s)
    db.session.flush()
    info_table = student.Info(student_id=s.id, **info)
    db.session.add(info_table)
    if father:
        f = student.Father(
            name=father[1],
            lastname=father[0],
            student_id=s.id
        )
        print(father[2] if len(father) > 2 else None)
        f.patronymic = father[2] if len(father) > 2 else None
        f.phone = father[3] if len(father) > 3 else None
        db.session.add(f)
    if mother:
        m = student.Mother(
            name=mother[1],
            lastname=mother[0],
            student_id=s.id
        )
        m.patronymic = mother[2] if len(mother) > 2 else None
        m.phone = mother[3] if len(mother) > 3 else None
        db.session.add(m)
    cases = student.Cases(student_id=s.id)
    db.session.add(cases)
    db.session.flush()

    if not cases.change():
        pass
db.session.commit()
#
#
# doc: Document = Document()
#
# # for style in doc.styles:
# #     if style.type != WD_STYLE_TYPE.TABLE:
# #         continue
# #     style: _TableStyle
# #     paragraph = doc.add_paragraph(style.name)
# #     table: Table = doc.add_table(rows=3, cols=6)
# #     table.style = style
# #     print(style)
# from app.models.student import Students
#
# # doc.add_heading('Добавление заголовка документа', level=2)
#
# style: _ParagraphStyle = doc.styles["Normal"]
# header_style: _ParagraphStyle = doc.styles["Header"]
#
# header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
# header_style.font.color.rgb = RGBColor(0, 0, 0)
# header_style.font.size = Pt(16)
# header_style.font.bold = True
#
#
# style.font.color.rgb = RGBColor(0, 0, 0)
# section: Section = doc.sections[0]
# header: Paragraph = section.header.paragraphs[0]
# header.text = "Сведения 11 Класс"
# header.style = doc.styles["Header"]
# section.orientation = WD_ORIENTATION.LANDSCAPE
# section.page_height, section.page_width = Cm(20), Cm(30)
# style.font.name = "Arial Narrow"
# style.font.bold = False
# style.font.size = Pt(10)
# table: Table = doc.add_table(rows=1, cols=8)
# fc: _Column = table.columns[0]
# table.style = "Table Grid"
# table_headers: _Row = table.rows[0]
# table_headers.height = 300000
#
# header_style: _ParagraphStyle = doc.styles.add_style("table_header", WD_STYLE_TYPE.PARAGRAPH)
# header_style.font.bold = True
# header_style.font.name = "Arial Narrow"
# header_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
# table.cell(0, 0).text = "№"
# table.cell(0, 1).text = "ФИО Ребенка"
# table.cell(0, 2).text = "Родители"
# table.cell(0, 1).width = table.cell(0, 2).width = 2400000
# table.cell(0, 3).text = "Телефоны"
# table.cell(0, 4).text = "Дата рождения"
# table.cell(0, 4).width = Cm(1)
# table.cell(0, 0).width = 300000
# for r in range(5):
#     table.cell(0, r).paragraphs[0].style = header_style
#
# for row_num, student in enumerate(Students.query.all(), 1):
#     student: Students
#     father = student.father.first()
#     mother = student.mother.first()
#     cells: tuple[_Cell] = table.add_row().cells
#     table.rows[row_num].height = 300000
#     cells[0].text = f"{row_num}"
#     cells[1].text = student.fullname()
#     parents = cells[2]
#     phones = cells[3]
#     cells[4].text = f"{student.birthdate_to_date().strftime('%d.%m.%Y')}"
#     cells[0].width = 300000
#     phones.text = parents.text = ""
#     if mother:
#         parents.text += mother.fullname() + ("\n" if father else "")
#         if mother.phone:
#             phones.text += mother.phone + "м" + ("\n" if father else "")
#     if father:
#         parents.text += father.fullname()
#         if father.phone:
#             phones.text += father.phone + "п"
#
#
# doc.save("test.docx")
